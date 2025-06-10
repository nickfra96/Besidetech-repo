import streamlit as st
import openai
import PyPDF2
import pandas as pd
import docx
import io
import json

# --- System Prompt per OpenAI ---
SYSTEM_PROMPT_MATCHER = """Sei un assistente AI esperto nella compilazione di risposte a criteri specifici basandoti sul contenuto di un documento fornito. Il tuo obiettivo è analizzare un documento per rispondere a una serie di criteri, come se stessi compilando una domanda o una valutazione dettagliata.

Riceverai:
1.  Una lista di criteri, ognuno con un "criterio_id" e una "descrizione_guida" che ne chiarisce il significato e il tipo di informazione richiesta.
2.  Il testo completo di un documento.

Per OGNI criterio fornito:
1.  Interpreta attentamente il "criterio_id" e la sua "descrizione_guida" per comprendere appieno cosa viene chiesto.
2.  Analizza l'intero testo del documento per trovare tutte le informazioni pertinenti, i dati, gli esempi, le descrizioni di processi o le affermazioni che dimostrano come il contenuto del documento (o il progetto/soggetto descritto nel documento) soddisfa, affronta o si riferisce a quel criterio.
3.  Sulla base delle informazioni trovate, formula una risposta completa, coerente e argomentata per quel criterio. Questa risposta dovrebbe spiegare in dettaglio come il documento/progetto risponde al criterio, utilizzando, citando (se utile, ad esempio "come menzionato nella sezione X" o "i dati a pagina Y indicano che...") e sintetizzando le informazioni estratte dal testo. Non limitarti a copiare porzioni di testo, ma elabora una risposta che sembri scritta da un esperto che sta compilando una domanda ufficiale e vuole essere esaustivo.
4.  Se il documento fornisce informazioni quantitative (numeri, percentuali, budget) o qualitative specifiche (descrizioni di qualità, metodologie, standard) relative al criterio, assicurati di includerle nella tua risposta per renderla più forte e basata su evidenze.
5.  Se il documento non contiene informazioni sufficienti o dirette per rispondere in modo esauriente a un criterio, indicalo chiaramente e onestamente nella risposta per quel criterio (es. "Il documento fornisce informazioni limitate riguardo a [aspetto specifico del criterio], menzionando solo che..." oppure "Non sono state trovate nel documento informazioni specifiche relative a [aspetto specifico del criterio], sebbene vengano trattati temi correlati come..."). Evita di inventare informazioni.

Restituisci i risultati come un elenco JSON di oggetti. Ogni oggetto nell'elenco deve corrispondere a uno dei criteri di input e contenere:
- "criterio_id": l'ID del criterio originale fornito.
- "descrizione_guida": la descrizione guida originale del criterio.
- "risposta_al_criterio_dal_documento": la risposta dettagliata, argomentata e formulata da te, basata sull'analisi approfondita del testo del documento in relazione al criterio.

Esempio di input dei criteri (parte del messaggio utente):
CRITERI A CUI RISPONDERE:
[
  {"criterio_id": "SOSTENIBILITA-AMBIENTALE", "descrizione_guida": "Descrivere in dettaglio le misure specifiche adottate e pianificate per garantire la sostenibilità ambientale del progetto, inclusi impatti, mitigazioni e monitoraggio."},
  {"criterio_id": "INNOVAZIONE-TECNOLOGICA", "descrizione_guida": "Illustrare il grado di innovazione tecnologica introdotto dal progetto, specificando le tecnologie utilizzate e il loro carattere innovativo rispetto allo stato dell'arte."},
  {"criterio_id": "IMPATTO-OCCUPAZIONALE", "descrizione_guida": "Valutare l'impatto occupazionale diretto e indiretto del progetto, fornendo stime numeriche se disponibili."}
]

Esempio di output JSON desiderato:
[
  {
    "criterio_id": "SOSTENIBILITA-AMBIENTALE",
    "descrizione_guida": "Descrivere in dettaglio le misure specifiche adottate e pianificate per garantire la sostenibilità ambientale del progetto, inclusi impatti, mitigazioni e monitoraggio.",
    "risposta_al_criterio_dal_documento": "Il progetto affronta la sostenibilità ambientale attraverso diverse misure chiave, come dettagliato nel documento. In primo luogo, si prevede l'installazione di un impianto fotovoltaico di 49,50 kW, che si aggiunge a quelli esistenti, con l'obiettivo di raggiungere una produzione di 200kW per ridurre significativamente l'impronta ambientale (come menzionato in relazione al criterio A1.1). Viene inoltre specificato che la riorganizzazione della sala ristorante, pur ampliandola, non comporta ulteriore consumo di suolo, in quanto si utilizzano spazi esistenti. La climatizzazione sarà ad alta efficienza energetica e la divisione della sala permetterà un consumo energetico mirato e ridotto. Il documento evidenzia anche l'uso di attrezzature CAM per il parco giochi, suggerendo un'attenzione ai Criteri Ambientali Minimi. Non vengono, tuttavia, fornite informazioni esplicite su piani di monitoraggio continuo degli impatti ambientali post-intervento o su specifiche strategie di mitigazione per altri potenziali impatti al di là del consumo energetico e di suolo."
  },
  {
    "criterio_id": "INNOVAZIONE-TECNOLOGICA",
    "descrizione_guida": "Illustrare il grado di innovazione tecnologica introdotto dal progetto, specificando le tecnologie utilizzate e il loro carattere innovativo rispetto allo stato dell'arte.",
    "risposta_al_criterio_dal_documento": "L'innovazione tecnologica del progetto si manifesta principalmente nell'adozione di soluzioni per l'efficienza energetica e l'accessibilità. L'impianto fotovoltaico da 49,50 kW rappresenta un potenziamento tecnologico per l'autonomia energetica. L'impianto di climatizzazione per la sala ristorante è descritto come 'ad alta efficienza e prestazione energetica con gestione digitale delle temperature' (D2.1), suggerendo l'uso di tecnologie moderne per il controllo ambientale. Un elemento innovativo significativo per l'accessibilità è l'installazione di un percorso tattile plantare integrato tipo Los-Vet-EVolution (LVE) per non vedenti, che migliora l'inclusività della struttura. La digitalizzazione è presente anche nella gestione delle temperature e nella programmazione delle sale polifunzionali. Il documento non confronta esplicitamente queste tecnologie con lo stato dell'arte del settore turistico, ma l'integrazione di queste soluzioni contribuisce a un innalzamento del livello tecnologico dell'offerta."
  },
  {
    "criterio_id": "IMPATTO-OCCUPAZIONALE",
    "descrizione_guida": "Valutare l'impatto occupazionale diretto e indiretto del progetto, fornendo stime numeriche se disponibili.",
    "risposta_al_criterio_dal_documento": "Il documento fornito non contiene informazioni esplicite o stime numeriche relative all'impatto occupazionale diretto o indiretto che deriverà dalla realizzazione degli interventi proposti. L'enfasi è posta sulla riqualificazione delle infrastrutture, l'incremento dei flussi turistici e la destagionalizzazione, che potrebbero implicitamente suggerire un mantenimento o un potenziale aumento dell'occupazione, ma non vengono forniti dati specifici a riguardo."
  }
]

Assicurati che l'output sia ESATTAMENTE una lista JSON valida di oggetti come specificato. La "risposta_al_criterio_dal_documento" deve essere una tua elaborazione analitica e argomentata basata sul testo, non una semplice estrazione di frasi.
"""

# --- Funzioni di Estrazione Testo ---
def extract_text_from_pdf(file_bytes):
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        text = ""
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text += page.extract_text() or ""
        return text
    except Exception as e:
        st.error(f"Errore durante la lettura del PDF: {e}")
        return None

def extract_text_from_excel(file_bytes):
    try:
        xls = pd.ExcelFile(io.BytesIO(file_bytes))
        text = ""
        for sheet_name in xls.sheet_names:
            df = xls.parse(sheet_name, header=None)
            for col in df.columns:
                text += df[col].astype(str).str.cat(sep=' ') + " "
            text += "\n"
        return text.strip()
    except Exception as e:
        st.error(f"Errore durante la lettura del file Excel: {e}")
        return None

def extract_text_from_docx(file_bytes):
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\t"
                text += "\n"
        return text
    except Exception as e:
        st.error(f"Errore durante la lettura del file DOCX: {e}")
        return None

# --- Funzione per caricare i criteri da JSON  ---
def load_criteria_from_json(json_file_bytes):
    try:
        data = json.load(io.BytesIO(json_file_bytes))
        transformed_criteria = []
        source_list = None
        if isinstance(data, list):
            source_list = data
        elif isinstance(data, dict) and len(data) == 1:
            key = next(iter(data))
            if isinstance(data[key], list):
                source_list = data[key]
            else:
                st.error(f"Il file JSON dei criteri ha una chiave radice '{key}', ma il suo valore non è una lista.")
                return None
        
        if source_list:
            for item in source_list:
                if isinstance(item, dict) and len(item) == 1:
                    criterio_id = next(iter(item))
                    descrizione = item[criterio_id]
                    transformed_criteria.append({
                        "criterio_id": str(criterio_id),
                        "descrizione_guida": str(descrizione)
                    })
                elif isinstance(item, dict) and "criterio_id" in item:
                    desc_guida = item.get("descrizione_guida", item.get("descrizione", ""))
                    transformed_criteria.append({
                        "criterio_id": str(item["criterio_id"]),
                        "descrizione_guida": str(desc_guida)
                    })
                else:
                    st.error(f"Formato dell'elemento non riconosciuto nella lista dei criteri: {item}.")
                    return None
            if not transformed_criteria and source_list:
                 st.warning("La lista dei criteri nel JSON sembra essere vuota o in un formato non gestito.")
                 return []
            return transformed_criteria
        else:
            st.error("Formato JSON dei criteri non supportato.")
            return None
    except json.JSONDecodeError:
        st.error("Errore nel decodificare il file JSON dei criteri.")
        return None
    except Exception as e:
        st.error(f"Errore imprevisto durante il caricamento dei criteri: {e}")
        return None

# --- Funzione per chiamare OpenAI ---
def get_matched_text_from_openai(criteria_list, document_text, api_key):
    if not document_text:
        st.warning("Il contenuto del documento sorgente è vuoto.")
        return []
    if not criteria_list:
        st.warning("La lista dei criteri è vuota o non caricata.")
        return []

    criteria_input_str = json.dumps(criteria_list, indent=2, ensure_ascii=False)
    user_prompt = f"""CRITERI A CUI RISPONDERE:
{criteria_input_str}

TESTO DEL DOCUMENTO DA ANALIZZARE:
{document_text}

Per favore, analizza il testo del documento e, per ciascun criterio fornito, formula una risposta dettagliata come specificato nelle istruzioni di sistema.
"""

    try:
        client = openai.OpenAI(api_key=api_key)
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT_MATCHER},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.1,
            response_format={"type": "json_object"}
        )
        
        json_string_response = response.choices[0].message.content
        
        try:
            data = json.loads(json_string_response)
            parsed_results = []

            if isinstance(data, list):
                parsed_results = data
            elif isinstance(data, dict):
                found_list = None
                for key in data:
                    if isinstance(data[key], list):
                        found_list = data[key]
                        st.info(f"Risultati trovati sotto la chiave '{key}' nel JSON di risposta.")
                        break
                if found_list is not None:
                    parsed_results = found_list
                else:
                    st.warning(f"La risposta JSON era un dizionario, ma non conteneva una lista riconoscibile: {json_string_response}")
                    return []
            else:
                st.warning(f"Formato JSON inatteso: {json_string_response}")
                return []

            validated_results = []
            # Il prompt ora chiede "risposta_al_criterio_dal_documento"
            for item in parsed_results:
                if isinstance(item, dict) and "criterio_id" in item and "risposta_al_criterio_dal_documento" in item:
                    validated_results.append(item)
                else:
                    st.warning(f"Elemento risultato ignorato: formato imprevisto o chiavi mancanti ('criterio_id', 'risposta_al_criterio_dal_documento'). Elemento: {item}")
            return validated_results

        except json.JSONDecodeError:
            st.error(f"Errore nel decodificare la risposta JSON da OpenAI: {json_string_response}")
            return []
        except Exception as e_parse:
            st.error(f"Errore durante il parsing della risposta: {e_parse}. Risposta: {json_string_response}")
            return []

    except openai.APIConnectionError as e:
        st.error(f"Errore di connessione all'API OpenAI: {e}")
    except openai.RateLimitError as e:
        st.error(f"Rate limit superato per l'API OpenAI: {e}")
    except openai.AuthenticationError as e:
        st.error(f"Errore di autenticazione API OpenAI: {e}. Controlla la tua API Key.")
    except openai.APIError as e:
        st.error(f"Errore generico API OpenAI: {e}")
        if "context_length_exceeded" in str(e).lower():
            st.warning("Il testo del documento e/o la lista dei criteri potrebbero essere troppo lunghi.")
    except Exception as e:
        st.error(f"Errore imprevisto durante la chiamata a OpenAI: {e}")
    return []


# --- Interfaccia Streamlit  ---
st.set_page_config(layout="wide", page_title="Besidetech AI")

st.title("Besidetech AI")
st.markdown("""
Carica un file JSON con i criteri (e le loro descrizioni guida) e un documento sorgente (PDF, Word, Excel). 
L'applicazione utilizzerà `gpt-4o-mini` di OpenAI per analizzare il documento e generare risposte argomentate per ciascun criterio, 
come se stesse compilando una domanda.
""")

col1, col2 = st.columns(2)

with col1:
    st.subheader("1. Carica i Criteri")
    uploaded_criteria_file = st.file_uploader("📂 Carica file JSON dei Criteri", type=["json"], help="Il JSON deve contenere una lista di oggetti, ognuno con 'criterio_id' e 'descrizione_guida' (o un formato trasformabile).")
    
    st.subheader("2. Carica il Documento Sorgente")
    uploaded_document_file = st.file_uploader("📄 Carica Documento (PDF, DOCX, XLSX)", type=["pdf", "docx", "xlsx", "xls"])

with col2:
    st.subheader("3. Configurazione API")
    api_key = st.text_input("🔑 Inserisci la tua API Key di OpenAI", type="password")

criteria_list = None
document_text = None
document_file_name = None

if uploaded_criteria_file:
    criteria_list = load_criteria_from_json(uploaded_criteria_file.getvalue())
    if criteria_list is not None:
        if criteria_list:
             st.success(f"File JSON dei criteri '{uploaded_criteria_file.name}' caricato. Trovati {len(criteria_list)} criteri.")
             with st.expander("Visualizza Criteri Caricati (formato per AI)"):
                st.json(criteria_list[:5]) 
                if len(criteria_list) > 5:
                    st.caption(f"... e altri {len(criteria_list)-5} criteri.")
        elif uploaded_criteria_file.size > 0 :
            st.warning(f"File JSON '{uploaded_criteria_file.name}' caricato, ma nessun criterio valido estratto.")
    # Errori gestiti da load_criteria_from_json

if uploaded_document_file:
    doc_bytes = uploaded_document_file.getvalue()
    document_file_name = uploaded_document_file.name
    
    with st.spinner(f"Estrazione del testo da `{document_file_name}`..."):
        if document_file_name.lower().endswith(".pdf"):
            document_text = extract_text_from_pdf(doc_bytes)
        elif document_file_name.lower().endswith((".xlsx", ".xls")):
            document_text = extract_text_from_excel(doc_bytes)
        elif document_file_name.lower().endswith(".docx"):
            document_text = extract_text_from_docx(doc_bytes)
    
    if document_text:
        st.success(f"Testo estratto da '{document_file_name}'. Lunghezza: {len(document_text)} caratteri.")
        with st.expander("Anteprima Testo Estratto (primi 500 caratteri)"):
            st.text_area("Anteprima", document_text[:500]+"...", height=150, disabled=True)
    elif uploaded_document_file:
        st.error(f"Non è stato possibile estrarre il testo da '{document_file_name}'.")


if st.button("✨ Analizza e Genera Risposte con OpenAI", disabled=(criteria_list is None or not criteria_list or not document_text or not api_key), help="Avvia l'analisi del documento e la generazione delle risposte ai criteri."):
    if not api_key:
        st.warning("Per favore, inserisci la tua API Key di OpenAI.")
    elif criteria_list is None or not criteria_list:
        st.warning("Per favore, carica un file JSON di criteri valido e non vuoto.")
    elif not document_text:
        st.warning("Per favore, carica un documento sorgente valido.")
    else:
        with st.spinner("Analisi del documento e generazione delle risposte con OpenAI in corso... Attendere prego."):
            generated_responses = get_matched_text_from_openai(criteria_list, document_text, api_key)
        
        st.subheader("Risposte Generate ai Criteri")
        if generated_responses:
            st.success(f"Generazione risposte completata. Elaborati {len(generated_responses)} criteri.")
            
            results_df = pd.DataFrame(generated_responses)
            
            # Configurazione colonne per st.data_editor
            # Il nome della colonna chiave è ora "risposta_al_criterio_dal_documento"
            column_config = {
                "criterio_id": st.column_config.TextColumn("ID Criterio", width="medium", help="L'ID del criterio fornito."),
                "descrizione_guida": st.column_config.TextColumn("Descrizione Guida (Input)", width="large", help="La descrizione guida originale del criterio."),
                "risposta_al_criterio_dal_documento": st.column_config.TextColumn("Risposta Generata dal Documento", width="extra_large", help="La risposta formulata dall'AI basata sull'analisi del documento.")
            }
            
            cols_to_display = {}
            if 'criterio_id' in results_df.columns:
                 cols_to_display['criterio_id'] = column_config['criterio_id']
            if 'descrizione_guida' in results_df.columns: 
                 cols_to_display['descrizione_guida'] = column_config['descrizione_guida']
            if 'risposta_al_criterio_dal_documento' in results_df.columns:
                 cols_to_display['risposta_al_criterio_dal_documento'] = column_config['risposta_al_criterio_dal_documento']


            st.data_editor(
                results_df,
                use_container_width=True,
                column_config=cols_to_display,
                hide_index=True,
                num_rows="dynamic",
                height=600 
            )
            
            json_output_responses = json.dumps(generated_responses, indent=2, ensure_ascii=False)
            st.download_button(
                label="📥 Scarica Risposte JSON",
                data=json_output_responses,
                file_name=f"risposte_criteri_{document_file_name.split('.')[0] if document_file_name else 'documento'}.json",
                mime="application/json"
            )
        else:
            st.info("Nessuna risposta generata da OpenAI o la risposta non era nel formato atteso. Verifica il System Prompt e la qualità del testo/criteri.")

