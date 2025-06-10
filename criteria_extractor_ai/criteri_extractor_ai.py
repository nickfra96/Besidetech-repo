import streamlit as st
import openai
import PyPDF2
import pandas as pd
import docx
import io
import json

# --- System Prompt per OpenAI  ---
SYSTEM_PROMPT = """Sei un assistente AI avanzato, specializzato nell'analisi semantica di documenti per identificare e estrarre i concetti chiave che fungono da criteri, requisiti, punti di valutazione, o sezioni tematiche principali, insieme alle loro descrizioni.

Il tuo compito è analizzare attentamente il testo fornito. Anche se non ci sono codici espliciti (come A1, B2.1), devi identificare le frasi o i paragrafi che stabiliscono regole, linee guida, specifiche, o argomenti centrali che potrebbero essere considerati "criteri" in un senso più ampio.

Per ogni "criterio" identificato:
1.  **Identificatore del Criterio (`criterio_id`):**
    * Se trovi un codice alfanumerico esplicito (es. A1, 1.2.3, Art. 5, CRITERIO X), usa quello.
    * Se il criterio è introdotto da un titolo di sezione o un'intestazione chiara e concisa (es. "Requisiti Tecnici", "Articolo 5: Protezione dei Dati"), usa quel titolo come ID.
    * Se non c'è un codice o un titolo ovvio, cerca di derivare un ID breve e significativo dalle prime parole chiave della descrizione del criterio (es. "Sicurezza_Dati_Personali", "Valutazione_Rischi_Operativi", "Procedura_Backup_Dati"). L'ID dovrebbe essere il più univoco e rappresentativo possibile. Evita ID troppo generici o eccessivamente lunghi.
    * Come ultima risorsa, se non è possibile derivare un ID significativo in altro modo, puoi usare un placeholder come "Criterio Inferito N" (dove N è un numero progressivo per distinguerli), ma privilegia sempre le opzioni precedenti.

2.  **Descrizione del Criterio (`descrizione`):**
    * Estrai il testo che definisce, spiega, o dettaglia di cosa tratta quel criterio, requisito o sezione tematica. Cattura il testo più rilevante e completo che ne costituisce la spiegazione principale o la definizione. Assicurati di includere l'intera frase o il paragrafo pertinente.

Restituisci i risultati come un elenco JSON di oggetti. Ogni oggetto deve contenere due chiavi:
- "criterio_id": una stringa con l'identificatore del criterio.
- "descrizione": una stringa con il testo descrittivo associato.

Esempio di output JSON desiderato (con alcuni ID inferiti o basati su titoli):
[
  {
    "criterio_id": "A1",
    "descrizione": "Questo è il testo che descrive il Criterio A1 e ne specifica i dettagli principali."
  },
  {
    "criterio_id": "Requisiti Hardware Minimi",
    "descrizione": "Il sistema deve essere compatibile con processori Intel i5 di ottava generazione o superiori e richiedere almeno 8GB di RAM e 256GB di spazio su disco SSD."
  },
  {
    "criterio_id": "Crittografia_Comunicazioni",
    "descrizione": "Tutte le comunicazioni tra il client e il server devono essere crittografate utilizzando TLS 1.2 o versioni successive, con algoritmi di cifratura approvati."
  },
  {
    "criterio_id": "Criterio Inferito 1",
    "descrizione": "Le password utente devono avere una lunghezza minima di 12 caratteri, includere lettere maiuscole, minuscole, numeri e simboli speciali, e devono essere cambiate ogni 90 giorni."
  }
]

Se nel testo non viene identificato alcun elemento che possa ragionevolmente essere interpretato come criterio, requisito o sezione tematica rilevante con una descrizione associata, restituisci una lista JSON vuota:
[]

Assicurati che l'output sia ESATTAMENTE una lista JSON valida di oggetti come specificato. Non includere spiegazioni, commenti o testo al di fuori della struttura JSON. Sii diligente nel trovare descrizioni complete e pertinenti. L'obiettivo è cogliere l'essenza di ogni punto chiave del documento che possa fungere da criterio o requisito. Evita di creare criteri da frasi troppo brevi o frammentarie se non rappresentano chiaramente un punto di valutazione o una regola.
"""

# --- Funzioni di Estrazione Testo  ---
def extract_text_from_pdf(file_bytes):
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        text = ""
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text += page.extract_text() or ""
        return text
    except Exception as e:
        st.error(f"Errore durante la lettura del PDF: {e}")
        return None

def extract_text_from_excel(file_bytes):
    try:
        xls = pd.ExcelFile(io.BytesIO(file_bytes))
        text = ""
        for sheet_name in xls.sheet_names:
            df = xls.parse(sheet_name, header=None)
            for col in df.columns:
                text += df[col].astype(str).str.cat(sep=' ') + " "
            text += "\n"
        return text.strip()
    except Exception as e:
        st.error(f"Errore durante la lettura del file Excel: {e}")
        return None

def extract_text_from_docx(file_bytes):
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\t"
                text += "\n"
        return text
    except Exception as e:
        st.error(f"Errore durante la lettura del file DOCX: {e}")
        return None

# --- Funzione per chiamare OpenAI ---
def get_criteria_from_openai(text_content, api_key):
    if not text_content:
        st.warning("Il contenuto del file è vuoto o non è stato possibile estrarlo.")
        return []
    try:
        client = openai.OpenAI(api_key=api_key)
        # Possiamo aumentare aumentare max_tokens in base a quanto vogliamo la risposta lunga
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": f"Ecco il testo del documento da cui estrarre i criteri e le loro descrizioni:\n\n{text_content}"}
            ],
            temperature=0.0, # Facendo cosi il modello lo rendo meno creativo e più affine ai dati che gli passo"
            response_format={"type": "json_object"}
            #max_tokens=4000 # 
        )
        
        json_string_response = response.choices[0].message.content
        
        try:
            data = json.loads(json_string_response)

            if isinstance(data, list):
                parsed_criteria = data
            elif isinstance(data, dict):
                found_list = None
                
                common_keys = ["criteri", "criteria", "results", "items", "data", "extracted_criteria"]
                for potential_key in common_keys:
                    if potential_key in data and isinstance(data[potential_key], list):
                        found_list = data[potential_key]
                        break
                if not found_list: 
                    for key_in_dict in data:
                        if isinstance(data[key_in_dict], list):
                            found_list = data[key_in_dict]
                            st.info(f"Trovata lista di criteri sotto la chiave generica '{key_in_dict}' nel JSON.")
                            break
                
                if found_list is not None:
                    parsed_criteria = found_list
                else:
                    st.warning(f"La risposta JSON era un dizionario, ma non conteneva una lista di criteri riconoscibile: {json_string_response}")
                    return []
            else:
                st.warning(f"Formato JSON inatteso. La risposta non è una lista né un dizionario contenente una lista di criteri: {json_string_response}")
                return []

            validated_criteria = []
            for item in parsed_criteria:
                if isinstance(item, dict) and "criterio_id" in item and "descrizione" in item:
                 
                    item["criterio_id"] = str(item["criterio_id"]).strip()
                    item["descrizione"] = str(item["descrizione"]).strip()
                    if item["criterio_id"] and item["descrizione"]: #  non siano stringhe vuote dopo lo strip
                         validated_criteria.append(item)
                    else:
                        st.warning(f"Elemento ignorato: 'criterio_id' o 'descrizione' vuoti dopo la pulizia. Originale: {item}")
                else:
                    st.warning(f"Elemento ignorato: formato imprevisto o chiavi mancanti ('criterio_id', 'descrizione') nell'oggetto JSON: {item}")
            return validated_criteria

        except json.JSONDecodeError:
            st.error(f"Errore nel decodificare la risposta JSON da OpenAI: {json_string_response}")
            st.info("Potrebbe essere necessario aggiustare il system prompt, controllare la risposta del modello, o il testo inviato potrebbe essere troppo lungo/complesso per un output JSON coerente con questo modello.")
            return []
        except Exception as e_parse:
            st.error(f"Errore durante il parsing della risposta strutturata: {e_parse}. Risposta grezza: {json_string_response}")
            return []

    except openai.APIConnectionError as e:
        st.error(f"Errore di connessione all'API OpenAI: {e}")
    except openai.RateLimitError as e:
        st.error(f"Rate limit superato per l'API OpenAI: {e}")
    except openai.AuthenticationError as e:
        st.error(f"Errore di autenticazione API OpenAI: {e}. Controlla la tua API Key.")
    except openai.APIError as e: 
        st.error(f"Errore API OpenAI: {e}")
        if "context_length_exceeded" in str(e).lower():
            st.warning("Il testo del documento potrebbe essere troppo lungo per il modello selezionato. Bisogna usaree un documento più corto o un modello con una finestra di contesto maggiore, se disponibile.")
    except Exception as e:
        st.error(f"Errore imprevisto durante la chiamata a OpenAI: {e}")
    return []

# --- Interfaccia Streamlit ---
st.set_page_config(layout="wide", page_title="Besidetech Extradtor Criteri")

st.title("📄 Besidetech Extractor Criteri From Files")
st.markdown("""
Carica un file (PDF, Excel, Word DOCX) per estrarre criteri, requisiti o sezioni tematiche chiave e le relative descrizioni, anche quando non esplicitamente codificati.
L'applicazione utilizza il modello `gpt-4o-mini` di OpenAI per interpretare il testo.
L'output sarà un JSON strutturato con "criterio_id" e "descrizione".
""")

api_key = st.text_input("🔑 Inserisci la tua API Key di OpenAI", type="password")
uploaded_file = st.file_uploader("📂 Carica il tuo documento", type=["pdf", "xlsx", "xls", "docx"])

if uploaded_file is not None:
    file_bytes = uploaded_file.getvalue()
    file_name = uploaded_file.name
    st.write(f"File caricato: `{file_name}`")

    text_content = None
    with st.spinner(f"Estrazione del testo da `{file_name}`..."):
        if file_name.lower().endswith(".pdf"):
            text_content = extract_text_from_pdf(file_bytes)
        elif file_name.lower().endswith((".xlsx", ".xls")):
            text_content = extract_text_from_excel(file_bytes)
        elif file_name.lower().endswith(".docx"):
            text_content = extract_text_from_docx(file_bytes)
        else:
            st.error("Formato file non supportato.")

    if text_content:
        max_preview_chars = 1000
        preview_text = text_content[:max_preview_chars]
        if len(text_content) > max_preview_chars:
            preview_text += "..."
        
        st.subheader(f"Anteprima del Testo Estratto (primi {max_preview_chars} caratteri circa)")
        st.text_area("Testo Estratto", preview_text, height=150, disabled=True, help=f"Lunghezza totale del testo estratto: {len(text_content)} caratteri.")

        if api_key:
            if st.button("Estrai Criteri e Descrizioni con OpenAI"):
                if len(text_content) < 50: # Controllo minimo sulla lunghezza del testo
                    st.warning("Il testo estratto sembra troppo corto per un'analisi significativa. Verifica il contenuto del file.")
                else:
                    with st.spinner("Analisi del testo con OpenAI in corso... Potrebbe richiedere qualche istante, specialmente per documenti lunghi."):
                        criteria_data = get_criteria_from_openai(text_content, api_key) 
                    
                    st.subheader("✅ Criteri Estratti con Descrizioni")
                    if criteria_data:
                        st.success(f"Trovati {len(criteria_data)} criteri/sezioni con descrizioni.")
                        
                        criteria_df = pd.DataFrame(criteria_data)
                        
                        st.data_editor(
                            criteria_df,
                            use_container_width=True,
                            column_config={
                                "criterio_id": st.column_config.TextColumn("ID Criterio/Sezione", width="medium", help="Identificatore del criterio (esplicito o inferito)."),
                                "descrizione": st.column_config.TextColumn("Descrizione", width="large", help="Testo che descrive il criterio/sezione.")
                            },
                            hide_index=True,
                            num_rows="dynamic" # Permette di vedere più righe se lo spazio lo consente
                        )
                        
                        json_output = json.dumps(criteria_data, indent=2, ensure_ascii=False)
                        st.download_button(
                            label="📥 Scarica JSON Strutturato",
                            data=json_output,
                            file_name=f"criteri_dettagliati_{file_name.split('.')[0]}.json",
                            mime="application/json"
                        )
                    else:
                        st.info("Nessun criterio (con descrizione) trovato nel documento secondo il modello OpenAI, o la risposta non era nel formato JSON atteso. Controllare anche il System Prompt o la qualità del testo estratto.")
        else:
            st.warning("Inserisci la tua API Key di OpenAI per procedere.")
    elif uploaded_file:
        st.error("Non è stato possibile estrarre il testo dal file. Controlla i messaggi di errore sopra.")

st.sidebar.markdown("---")
